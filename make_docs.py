# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.insert(0, bidi)

def set_align_right(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def add_rtl_para(doc, text, bold=False, size=11, color=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    set_rtl(p)
    set_align_right(p)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Arial'
    if color:
        run.font.color.rgb = RGBColor(*color)
    # RTL run
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)
    return p

def add_heading(doc, text, level=1, color=(26,74,46)):
    p = doc.add_heading(text, level=level)
    set_rtl(p)
    set_align_right(p)
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Arial'
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)
    return p

def add_space(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return p

def set_doc_rtl(doc):
    sectPr = doc.sections[0]._sectPr
    bidi = OxmlElement('w:bidi')
    sectPr.insert(0, bidi)

# ── DOCUMENT 1: EULA ──────────────────────────────────────────────────────────
def make_eula():
    doc = Document()
    set_doc_rtl(doc)

    # Margins
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Title
    p = doc.add_paragraph()
    set_rtl(p); set_align_right(p)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("הסכם רישיון משתמש קצה (EULA)")
    r.bold = True; r.font.size = Pt(20); r.font.name = 'Arial'
    r.font.color.rgb = RGBColor(26, 74, 46)

    p2 = doc.add_paragraph()
    set_rtl(p2); set_align_right(p2)
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run("תוכנת Bridge Teacher")
    r2.bold = True; r2.font.size = Pt(14); r2.font.name = 'Arial'
    r2.font.color.rgb = RGBColor(85, 85, 85)

    add_rtl_para(doc, "קרא הסכם זה בעיון לפני השימוש בתוכנה. שימוש בתוכנה מהווה הסכמה לתנאים המפורטים להלן.",
                 color=(85,85,85), size=10)

    add_space(doc)
    add_heading(doc, "1. הגדרות")
    add_rtl_para(doc, "בהסכם זה:")
    add_bullet(doc, '"התוכנה" — מוצר Bridge Teacher, לרבות כל עדכון או גרסה עתידית שלו.')
    add_bullet(doc, '"המפתח" — יצחק תילוביץ, בעל זכויות היוצרים הבלעדיות בתוכנה.')
    add_bullet(doc, '"המשתמש" — האדם או הגוף שרכש רישיון לשימוש בתוכנה.')

    add_space(doc)
    add_heading(doc, "2. מתן רישיון")
    add_rtl_para(doc, "המפתח מעניק למשתמש רישיון אישי, מוגבל, בלתי-עביר ולא-בלעדי להתקין ולהשתמש בתוכנה על מחשב אחד בבעלות המשתמש.")
    add_rtl_para(doc, "הרישיון ניתן למשתמש יחיד בלבד ואינו ניתן להעברה לאחרים ללא אישור מפורש בכתב מהמפתח.")

    add_space(doc)
    add_heading(doc, "3. הגבלות שימוש")
    add_rtl_para(doc, "המשתמש אינו רשאי:", bold=True)
    add_bullet(doc, "להעתיק, לשכפל או להפיץ את התוכנה בכל צורה שהיא.")
    add_bullet(doc, "למכור, להשכיר, להשאיל או להעביר את הרישיון לצד שלישי.")
    add_bullet(doc, "לבצע הנדסה לאחור (Reverse Engineering), פירוק (Decompile) או פיצוח של הקוד.")
    add_bullet(doc, "לשנות, לתרגם או ליצור עבודות נגזרות מהתוכנה.")
    add_bullet(doc, "להסיר או לשנות הודעות זכויות יוצרים המופיעות בתוכנה.")

    add_space(doc)
    add_heading(doc, "4. זכויות יוצרים וקניין רוחני")
    add_rtl_para(doc, "התוכנה, לרבות כל הקוד, העיצוב, הלוגיקה וחומרי העזר, הינה רכושו הבלעדי של יצחק תילוביץ. כל הזכויות שמורות.")
    add_rtl_para(doc, "הסכם זה אינו מעביר למשתמש כל זכות קניין רוחני בתוכנה.")

    add_space(doc)
    add_heading(doc, "5. הגבלת אחריות")
    add_rtl_para(doc, 'התוכנה מסופקת "AS IS" — כפי שהיא — ללא כל אחריות מפורשת או משתמעת.')
    add_rtl_para(doc, "המפתח אינו אחראי לכל נזק ישיר, עקיף, מקרי או תוצאתי הנובע משימוש בתוכנה.")

    add_space(doc)
    add_heading(doc, "6. סיום ההסכם")
    add_rtl_para(doc, "הפרת תנאי כלשהו תגרור ביטול הרישיון אוטומטית וללא הודעה מוקדמת.")
    add_rtl_para(doc, "עם סיום ההסכם, המשתמש מחויב להסיר את התוכנה ממחשבו.")

    add_space(doc)
    add_heading(doc, "7. חוק שחל")
    add_rtl_para(doc, "הסכם זה כפוף לחוקי מדינת ישראל. כל סכסוך יידון בבתי המשפט המוסמכים בישראל.")

    add_space(doc)
    add_heading(doc, "8. אישור וחתימה")
    add_rtl_para(doc, "בחתימתי על הסכם זה אני מאשר כי קראתי, הבנתי והסכמתי לכל תנאיו.")
    add_space(doc)
    add_space(doc)
    add_rtl_para(doc, "תאריך: _______________________", size=11)
    add_space(doc)
    add_rtl_para(doc, "שם מלא: _______________________", size=11)
    add_space(doc)
    add_rtl_para(doc, "חתימה: _______________________", size=11)

    doc.save(r'D:\bridge-teacher1\הסכם_רישיון.docx')
    print("✓ הסכם_רישיון.docx")

# ── DOCUMENT 2: User Manual ───────────────────────────────────────────────────
def make_manual():
    doc = Document()
    set_doc_rtl(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Title
    p = doc.add_paragraph()
    set_rtl(p); set_align_right(p)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Bridge Teacher — מדריך למשתמש")
    r.bold = True; r.font.size = Pt(22); r.font.name = 'Arial'
    r.font.color.rgb = RGBColor(26, 74, 46)

    p2 = doc.add_paragraph()
    set_rtl(p2); set_align_right(p2)
    p2.paragraph_format.space_after = Pt(16)
    r2 = p2.add_run("עזרה למורה ברידג'")
    r2.italic = True; r2.font.size = Pt(14); r2.font.name = 'Arial'
    r2.font.color.rgb = RGBColor(85, 85, 85)

    add_heading(doc, "מה התוכנה עושה?")
    add_rtl_para(doc, "Bridge Teacher היא תוכנת שולחן עבודה למורי ברידג'. התוכנה מאפשרת יצירת לוחות אימון מותאמים אישית לפי שיטת Acol — 5 קלפים במיגור, 1NT חזק 15-17.")
    add_rtl_para(doc, "המורה מגדיר אילוצים לכל שחקן והתוכנה מייצרת לוחות עם הכרזות אוטומטיות מלאות, מוכנים להדפסה.")

    add_space(doc)
    add_heading(doc, "התקנה והפעלה")
    add_bullet(doc, "לחץ פעמיים על BridgeTeacher.exe")
    add_bullet(doc, "לא נדרשת התקנה — קובץ אחד, מריץ ישירות")
    add_bullet(doc, "דרישות מערכת: Windows 10 / Windows 11")
    add_bullet(doc, "גודל קובץ: כ-10 MB")

    add_space(doc)
    add_heading(doc, "הגדרות שיטה")
    add_rtl_para(doc, "בחלק העליון של המסך ניתן לשנות את פרמטרי השיטה:")
    add_bullet(doc, "מיגור: 4 קלפים / 5 קלפים (ברירת מחדל: 5 קלפים)")
    add_bullet(doc, "1NT: 12-14 / 15-17 (ברירת מחדל: 15-17)")
    add_bullet(doc, "מינור: 2 קלפים / 3 קלפים (ברירת מחדל: 3 קלפים)")

    add_space(doc)
    add_heading(doc, "הגדרת שחקנים")
    add_rtl_para(doc, "המסך מכיל 4 פאנלים — צפון, מזרח, דרום, מערב. כל שחקן יכול להיות:")
    add_space(doc)

    add_heading(doc, "חופשי", level=2, color=(85,85,85))
    add_rtl_para(doc, "ללא אילוצים — התוכנה תחלק יד אקראית.")

    add_heading(doc, "פותח", level=2, color=(85,85,85))
    add_bullet(doc, "כוח: חלש 12-14 / בינוני 15-17 / חזק 18-21")
    add_bullet(doc, "סוג פתיחה: מיגור (♥/♠) / מינור (♣/♦) / 1NT")

    add_heading(doc, "משיב", level=2, color=(85,85,85))
    add_rtl_para(doc, "לאחר הגדרת פותח, השחקן השותף הופך למשיב אוטומטית.")
    add_bullet(doc, "כוח משיב: חלש 6-9 / בינוני 10-12 / חזק 13+")
    add_bullet(doc, "סוג תשובה: תמיכה / 1NT / 1♠ / סטיימן / טרנספר ♥/♠ / 2NT")

    add_space(doc)
    add_heading(doc, "יצירת לוחות")
    add_bullet(doc, "בחר מספר לוחות (1 עד 36)")
    add_bullet(doc, "בחר דילר (אופציונלי: צפון / מזרח / דרום / מערב / ללא)")
    add_bullet(doc, 'לחץ "⚡ צור לוחות"')
    add_bullet(doc, "הלוחות נפתחים בחלון הדפסה נפרד")

    add_space(doc)
    add_heading(doc, "הדפסה")
    add_bullet(doc, "4 לוחות לעמוד, פורמט A4 לאורך")
    add_bullet(doc, "כל לוח כולל: פריסת 4 ידיים + מצפן + טבלת ניקוד ריקה")
    add_bullet(doc, 'לחץ "הדפס" לשליחה למדפסת')

    add_space(doc)
    add_heading(doc, "שיעורים מוכנים")
    add_rtl_para(doc, 'כפתור "שיעורים" מציג קבוצות ידיים מוכנות מראש לפי נושא:')
    add_bullet(doc, "1NT — פס, סטיימן גיים/הזמנה, טרנספר ♥/♠, 3NT")
    add_bullet(doc, "מיגורים — 1♥/1♠ עם כל סוגי התשובות")
    add_bullet(doc, "מינורים — 1♣/1♦ עם כל סוגי התשובות")
    add_bullet(doc, "סלאם — Gerber, RKCB, הזמנת סלאם")
    add_bullet(doc, "פתיחות 2 חלשות — 2♥/2♠")
    add_bullet(doc, "2♣ חזקה — תשובה שלילית/חיובית")
    add_bullet(doc, "2NT (20-22) — סטיימן, טרנספר, 3NT")

    add_space(doc)
    add_heading(doc, "כפתורים נוספים")
    add_bullet(doc, '"אפס" — מנקה את כל ההגדרות ומחזיר לברירת מחדל')
    add_bullet(doc, '"הכרזות ✓/✗" — הצג/הסתר הכרזות על הלוחות')

    add_space(doc)
    add_heading(doc, "יצירת קשר ותמיכה")
    add_rtl_para(doc, "לשאלות, הצעות שיפור ותמיכה טכנית:")
    add_rtl_para(doc, "יצחק תילוביץ — מורה ברידג'", bold=True)

    add_space(doc)
    p = doc.add_paragraph()
    set_rtl(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Bridge Teacher © יצחק תילוביץ — כל הזכויות שמורות")
    r.italic = True; r.font.size = Pt(9); r.font.name = 'Arial'
    r.font.color.rgb = RGBColor(136, 136, 136)

    doc.save(r'D:\bridge-teacher1\מדריך_למשתמש.docx')
    print("✓ מדריך_למשתמש.docx")

make_eula()
make_manual()
